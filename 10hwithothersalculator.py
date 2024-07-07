# -*- coding: utf-8 -*-
"""10hwithothersalculator.ipynb

Automatically generated by Colab.

Original file is located at
    https://colab.research.google.com/drive/1GieRx_5qIeVnoVeMqD__L-wt1N_ZphcO
"""

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from scipy.optimize import minimize
import zipfile
import os

try:
    import pandas as pd
    import numpy as np
    import matplotlib.pyplot as plt
    from docx import Document
    from docx.shared import Inches
    from scipy.optimize import minimize
except ImportError:
    !pip install pandas numpy matplotlib python-docx scipy

def calculate_h_index(y_data):
    sorted_cites = np.sort(y_data)[::-1]
    h_index = np.max(np.where(sorted_cites >= np.arange(1, len(sorted_cites) + 1))[0]) + 1
    return h_index

def calculate_abc(M, N, h):
    a = (M * h**2) / (M * N - (M + N) * h)
    b = (M * N * (M - h) * (N - h)) * (h / (M * N - (M + N) * h))**2
    c = (N * h**2) / (M * N - (M + N) * h)
    return a, b, c

def calculate_abc_h_formula(M, N, h):
    a = -(-M*N + M*h**2 - M*h + M + N*h - h**2)/(-M*N + M*h + N*h - 2*h + 1)
    c = -(-M*N + M*h + N*h**2 - N*h + N - h**2)/(-M*N + M*h + N*h - 2*h + 1)
    b = (-M + h)*(M - 1)*(-N + h)*(N - 1)*(h - 1)**2/(-M*N + M*h + N*h - 2*h + 1)**2
    return a, b, c

def calculate_rmse(y_true, y_pred):
    return np.sqrt(np.mean((y_true - y_pred) ** 2))

def optimize_abc(x_data, y_data):
    def objective(params):
        a, b, c = params
        fitted_values = (b / (x_data + c)) - a
        return calculate_rmse(y_data, fitted_values)

    initial_guess = [0.1, 100, 72]
    result = minimize(objective, initial_guess, method='Nelder-Mead')
    return result.x

def optimize_power_law(x_data, y_data):
    def objective(params):
        C, lam = params
        fitted_values = C / (x_data ** lam)
        return calculate_rmse(y_data, fitted_values)

    initial_guess = [np.max(y_data), 1]
    result = minimize(objective, initial_guess, method='Nelder-Mead')
    return result.x

# Upload Excel file
uploaded_file = 'your_uploaded_file.xlsx'

xls = pd.ExcelFile(uploaded_file)

doc = Document()
doc.add_heading('Plots for Each Sheet', 0)

for sheet_name in xls.sheet_names:
    df = pd.read_excel(xls, sheet_name=sheet_name)
    x_data = df['Serial Number'].values
    y_data = df['ECC'].values

    M = np.max(y_data)
    N = len(y_data)
    h = calculate_h_index(y_data)

    # Original Model
    a, b, c = calculate_abc(M, N, h)
    fitted_values = (b / (x_data + c)) - a
    rmse_original = calculate_rmse(y_data, fitted_values)

    # Optimized Model
    a_opt, b_opt, c_opt = optimize_abc(x_data, y_data)
    fitted_values_opt = (b_opt / (x_data + c_opt)) - a_opt
    rmse_eqn2 = calculate_rmse(y_data, fitted_values_opt)

    # Power Law Model
    C_opt, lambda_opt = optimize_power_law(x_data, y_data)
    fitted_values_power_law = C_opt / (x_data ** lambda_opt)
    rmse_power_law = calculate_rmse(y_data, fitted_values_power_law)

    # h-formula Model
    a_h, b_h, c_h = calculate_abc_h_formula(M, N, h)
    fitted_values_h = (b_h / (x_data + c_h)) - a_h
    rmse_h = calculate_rmse(y_data, fitted_values_h)

    # Plotting
    plt.figure(figsize=(10, 6))
    plt.scatter(x_data, y_data, label='Original Data')
    plt.plot(x_data, fitted_values, color='red', linestyle=':', label='Fitted Values (Model 1)')
    plt.plot(x_data, fitted_values_opt, color='blue', linestyle='--', label='Optimized Fitted Values (Model 2)')
    plt.plot(x_data, fitted_values_power_law, color='green', linestyle='-', label='Fitted Values (Power Law)')
    plt.plot(x_data, fitted_values_h, color='purple', linestyle='-.', label='Fitted Values (H Formula)')

    plt.text(0.5, 0.8, f'Model 1 - a: {a:.2f}, b: {b:.2f}, c: {c:.2f}', transform=plt.gca().transAxes, color='red')
    plt.text(0.5, 0.75, f'Model 2 - a: {a_opt:.2f}, b: {b_opt:.2f}, c: {c_opt:.2f}, RMSE: {rmse_eqn2:.2f}', transform=plt.gca().transAxes, color='blue')
    plt.text(0.5, 0.7, f'Power Law - C: {C_opt:.2f}, λ: {lambda_opt:.2f}, RMSE: {rmse_power_law:.2f}', transform=plt.gca().transAxes, color='green')
    plt.text(0.5, 0.65, f'H Formula - a: {a_h:.2f}, b: {b_h:.2f}, c: {c_h:.2f}, RMSE: {rmse_h:.2f}', transform=plt.gca().transAxes, color='purple')

    plt.xlabel('Serial Number')
    plt.ylabel('Cites')
    plt.title(f'Plot for {sheet_name}')
    plt.legend()

    plot_filename = f'{sheet_name}_plot.png'
    plt.savefig(plot_filename)
    plt.close()

    # Add plots and details to the Word document
    doc.add_heading(sheet_name, level=1)
    doc.add_picture(plot_filename, width=Inches(6))
    doc.add_paragraph(f'Model 1 - a: {a:.2f}, b: {b:.2f}, c: {c:.2f}, RMSE: {rmse_original:.2f}')
    doc.add_paragraph(f'Model 2 - a: {a_opt:.2f}, b: {b_opt:.2f}, c: {c_opt:.2f}, RMSE: {rmse_eqn2:.2f}')
    doc.add_paragraph(f'Power Law - C: {C_opt:.2f}, λ: {lambda_opt:.2f}, RMSE: {rmse_power_law:.2f}')
    doc.add_paragraph(f'H Formula - a: {a_h:.2f}, b: {b_h:.2f}, c: {c_h:.2f}, RMSE: {rmse_h:.2f}')

# Save the Word document
doc_filename = 'plots_and_results.docx'
doc.save(doc_filename)

# Download the Word document
from google.colab import files
files.download(doc_filename)