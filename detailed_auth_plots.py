# -*- coding: utf-8 -*-
"""DETAILED AUTH PLOTS.ipynb

Automatically generated by Colab.

Original file is located at
    https://colab.research.google.com/drive/1_e8AeJSzu1GC19OgdCw4JfJ0RP7f8Pls
"""

!pip install pandas numpy matplotlib python-docx openpyxl scipy

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from google.colab import files
from scipy.optimize import minimize

uploaded = files.upload()

file_name = list(uploaded.keys())[0]
xls = pd.ExcelFile(file_name)

doc = Document()
doc.add_heading('Plots for Each Sheet', 0)

original_model_df = pd.DataFrame(columns=['Sheet', 'M', 'N', 'h', 'a', 'b', 'c', 'RMSE'])
optimized_model_df = pd.DataFrame(columns=['Sheet', 'a_opt', 'b_opt', 'c_opt', 'RMSEEQN2'])
power_law_model_df = pd.DataFrame(columns=['Sheet', 'C_opt', 'lambda_opt', 'RMSE_PowerLaw'])
h_formula_model_df = pd.DataFrame(columns=['Sheet', 'a_h', 'b_h', 'c_h', 'RMSE'])

def calculate_h_index(y_data):
    sorted_cites = np.sort(y_data)[::-1]
    h_index = np.max(np.where(sorted_cites >= np.arange(1, len(sorted_cites) + 1))[0]) + 1
    return h_index

def calculate_abc(M, N, h):
    a = (M * h**2) / (M * N - (M + N) * h)
    b = (M * N * (M - h) * (N - h)) * (h / (M * N - (M + N) * h))**2
    c = (N * h**2) / (M * N - (M + N) * h)
    return a, b, c

def calculate_abc_h_formula(M, N, h):
    a = -(-M*N + M*h**2 - M*h + M + N*h - h**2)/(-M*N + M*h + N*h - 2*h + 1)
    c = -(-M*N + M*h + N*h**2 - N*h + N - h**2)/(-M*N + M*h + N*h - 2*h + 1)
    b = (-M + h)*(M - 1)*(-N + h)*(N - 1)*(h - 1)**2/(-M*N + M*h + N*h - 2*h + 1)**2
    return a, b, c

def calculate_rmse(y_true, y_pred):
    return np.sqrt(np.mean((y_true - y_pred) ** 2))

def optimize_abc(x_data, y_data):
    def objective(params):
        a, b, c = params
        fitted_values = (b / (x_data + c)) - a
        return calculate_rmse(y_data, fitted_values)

    initial_guess = [0.1, 100, 72]
    result = minimize(objective, initial_guess, method='Nelder-Mead')
    return result.x

def optimize_power_law(x_data, y_data):
    def objective(params):
        C, lam = params
        fitted_values = C / (x_data ** lam)
        return calculate_rmse(y_data, fitted_values)

    initial_guess = [M, 1]
    result = minimize(objective, initial_guess, method='Nelder-Mead')
    return result.x

for sheet_name in xls.sheet_names:
    df = pd.read_excel(xls, sheet_name=sheet_name)
    x_data = df['Serial Number'].values
    y_data = df['ECC'].values

    M = np.max(y_data)
    N = len(y_data)
    h = calculate_h_index(y_data)

    a, b, c = calculate_abc(M, N, h)
    fitted_values = (b / (x_data + c)) - a
    rmse_original = calculate_rmse(y_data, fitted_values)

    a_opt, b_opt, c_opt = optimize_abc(x_data, y_data)
    fitted_values_opt = (b_opt / (x_data + c_opt)) - a_opt
    rmse_eqn2 = calculate_rmse(y_data, fitted_values_opt)

    C_opt, lambda_opt = optimize_power_law(x_data, y_data)
    fitted_values_power_law = C_opt / (x_data ** lambda_opt)
    rmse_power_law = calculate_rmse(y_data, fitted_values_power_law)

    a_h, b_h, c_h = calculate_abc_h_formula(M, N, h)
    fitted_values_h = (b_h / (x_data + c_h)) - a_h
    rmse_h = calculate_rmse(y_data, fitted_values_h)

    original_row = pd.DataFrame({
        'Sheet': [sheet_name],
        'M': [M],
        'N': [N],
        'h': [h],
        'a': [a],
        'b': [b],
        'c': [c],
        'RMSE': [rmse_original]
    })
    original_model_df = pd.concat([original_model_df, original_row], ignore_index=True)

    optimized_row = pd.DataFrame({
        'Sheet': [sheet_name],
        'a_opt': [a_opt],
        'b_opt': [b_opt],
        'c_opt': [c_opt],
        'RMSEEQN2': [rmse_eqn2]
    })
    optimized_model_df = pd.concat([optimized_model_df, optimized_row], ignore_index=True)

    power_law_row = pd.DataFrame({
        'Sheet': [sheet_name],
        'C_opt': [C_opt],
        'lambda_opt': [lambda_opt],
        'RMSE_PowerLaw': [rmse_power_law]
    })
    power_law_model_df = pd.concat([power_law_model_df, power_law_row], ignore_index=True)

    h_formula_row = pd.DataFrame({
        'Sheet': [sheet_name],
        'a_h': [a_h],
        'b_h': [b_h],
        'c_h': [c_h],
        'RMSE': [rmse_h]
    })
    h_formula_model_df = pd.concat([h_formula_model_df, h_formula_row], ignore_index=True)

    plt.figure(figsize=(10, 6))
    plt.scatter(x_data, y_data, label='Original Data')
    plt.plot(x_data, fitted_values, color='red', linestyle=':', label='Fitted Values (Model 1)')
    plt.plot(x_data, fitted_values_opt, color='blue', linestyle='--', label='Optimized Fitted Values (Model 2)')
    plt.plot(x_data, fitted_values_power_law, color='green', linestyle='-', label='Fitted Values (Power Law)')

    plt.text(0.5, 0.8, f'Model 1 - a: {a:.2f}, b: {b:.2f}, c: {c:.2f}', transform=plt.gca().transAxes, color='red')
    plt.text(0.5, 0.75, f'Model 2 - a: {a_opt:.2f}, b: {b_opt:.2f}, c: {c_opt:.2f}, RMSE: {rmse_eqn2:.2f}', transform=plt.gca().transAxes, color='blue')
    plt.text(0.5, 0.7, f'Power Law - C: {C_opt:.2f}, λ: {lambda_opt:.2f}, RMSE: {rmse_power_law:.2f}', transform=plt.gca().transAxes, color='green')

    plt.xlabel('Serial Number')
    plt.ylabel('Cites')
    plt.title(f'Plot for {sheet_name}')
    plt.legend()

    plot_filename = f'{sheet_name}.png'
    plt.savefig(plot_filename)
    plt.close()

    doc.add_heading(sheet_name, level=1)
    doc.add_picture(plot_filename, width=Inches(6))
    doc.add_paragraph(f'Model 1 - a: {a:.2f}, b: {b:.2f}, c: {c:.2f}')
    doc.add_paragraph(f'Model 2 - a: {a_opt:.2f}, b: {b_opt:.2f}, c: {c_opt:.2f}, RMSE: {rmse_eqn2:.2f}')
    doc.add_paragraph(f'Power Law - C: {C_opt:.2f}, λ: {lambda_opt:.2f}, RMSE: {rmse_power_law:.2f}')

doc_filename = 'fitted_citation_analysis.docx'
doc.save(doc_filename)

results_filename = 'results.xlsx'
with pd.ExcelWriter(results_filename) as writer:
    original_model_df.to_excel(writer, sheet_name='Original Model', index=False)
    optimized_model_df.to_excel(writer, sheet_name='Optimized Model', index=False)
    power_law_model_df.to_excel(writer, sheet_name='Power Law Model', index=False)
    h_formula_model_df.to_excel(writer, sheet_name='11h', index=False)

files.download(doc_filename)
files.download(results_filename)