import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from scipy.stats import linregress
from tkinter import Tk, filedialog
from docx import Document
from docx.shared import Inches
import os

def calculate_h_index(citations):
    citations_sorted = np.sort(citations)[::-1]
    h_index = np.sum(citations_sorted >= np.arange(1, len(citations_sorted) + 1))
    return h_index


def calculate_g_index_alt(citations):
    citations_sorted = np.sort(citations)[::-1]
    g_index = 0
    total_citations = 0

    for i, citation_count in enumerate(citations_sorted, start=1):
        total_citations += citation_count
        if total_citations >= i**2:
            g_index = i

    return g_index

def calculate_new_g_index(h_index, M):
    e = 2.71828
    new_g_index = h_index * (np.sqrt(np.log((4 * M) / (e * h_index))))
    return new_g_index

def calculate_a_index(h_index, citations):
    h_papers = citations[citations >= h_index]
    a_index = np.mean(h_papers)
    return a_index

def calculate_r_index(h_index, citations):
    h_papers = citations[citations >= h_index]
    r_index = np.sqrt(np.sum(h_papers))
    return r_index

def calculate_hg_index(h_index, g_index):
    hg_index = np.sqrt(h_index * g_index)
    return hg_index

def calculate_new_a_index(M, h_index):
    new_a_index = ((M * h_index) / (M - h_index)) * np.log(M / h_index)
    return new_a_index

def calculate_e_index(h_index, citations):
    h_papers = citations[citations >= h_index]
    excess_citations = np.sum(h_papers) - h_index**2
    e_index = np.sqrt(excess_citations)
    return e_index

def calculate_new_e_index(M, h_index):
    new_e_index_squared = ((M * h_index**2) / (M - h_index)) * np.log(M / h_index) - h_index**2
    new_e_index = np.sqrt(new_e_index_squared)
    return new_e_index

def calculate_new_r_index(M, h_index):
    new_r_index = h_index * np.sqrt((M / (M - h_index)) * np.log(M / h_index))
    return new_r_index

def calculate_new_hg_index(h_index, M):
    e = 2.71828
    new_hg_index = h_index * (np.log((4 * M) / (e * h_index)))**(1/4)
    return new_hg_index

def calculate_h_prime(h_index, citations):
    sorted_citations = np.sort(citations)[::-1]
    Fhead = np.sum(sorted_citations[:h_index])
    Ftail = np.sum(sorted_citations[h_index:])
    h_prime = np.sqrt(Fhead / Ftail) * h_index
    return h_prime

def calculate_new_h_prime(M, h_index, N):
    numerator = (M / (M - h_index)) * (np.log(M / h_index)) - 1
    denominator = (N / (N - h_index)) * (np.log(N / h_index)) - 1
    new_h_prime = np.sqrt(numerator / denominator) * h_index
    return new_h_prime

def calculate_h2_index(citations):
    citations_sorted = np.sort(citations)[::-1]
    h2_index = 0

    for i in range(1, len(citations_sorted) + 1):
        if citations_sorted[i-1] >= i**2:
            h2_index = i
        else:
            break

    return h2_index

def calculate_F1(M, N, h_index):
    numerator1 = M * (N - 1) * h_index * (h_index - 1)
    denominator1 = M * N - (M + N) * h_index + h_index
    
    term1 = numerator1 / denominator1
    
    numerator2 = (M - h_index) * (N - h_index)
    denominator2 = M * N - (M + N) * h_index + h_index
    log_term = np.log(N / ((N - 1) * h_index * (h_index - 1)))
    
    term2 = numerator2 / denominator2 * log_term * (M * N - (M + N) * h_index+ h_index) - 1
    
    F1 = term1 * term2
    return F1

def calculate_F2(M, N, h_index):
    numerator1 = M * (N - 1) * h_index * (h_index - 1)
    denominator1 = M * N - (M + N) * h_index + h_index
    
    term1 = numerator1 / denominator1
    
    log_term = np.log(N / np.e)
    
    numerator2 = M * N - (M + N) * h_index + h_index
    denominator2 = (N - 1) * h_index * (h_index - 1)
    
    term2 = numerator2 / denominator2
    
    F2 = term1 * log_term * term2
    return F2

def process_excel_files(file_paths):
    result_rows = []

    for file_path in file_paths:
        xl = pd.ExcelFile(file_path)
        
        for sheet_name in xl.sheet_names:
            df = xl.parse(sheet_name)
            citations = df['Cites'].values
            sum_of_cites = np.sum(citations)  # Calculate the sum of citations

            M = np.max(citations)
            N = len(citations)
            h_index = calculate_h_index(citations)
            g_index = calculate_g_index_alt(citations)
            new_g_index = calculate_new_g_index(h_index, M)
            a_index = calculate_a_index(h_index, citations)
            r_index = calculate_r_index(h_index, citations)
            hg_index = calculate_hg_index(h_index, g_index)
            new_r_index = calculate_new_r_index(M, h_index)
            new_hg_index = calculate_new_hg_index(h_index, M)
            new_a_index = calculate_new_a_index(M, h_index)
            e_index = calculate_e_index(h_index, citations)
            new_e_index = calculate_new_e_index(M, h_index)
            h_prime = calculate_h_prime(h_index, citations)
            new_h_prime = calculate_new_h_prime(M, h_index, N)
            h2_index = calculate_h2_index(citations)
            F1 = calculate_F1(M, N, h_index)
            F2 = calculate_F2(M, N, h_index)
            result_row = {
                'File': file_path,
                'Sheet': sheet_name,
                'M': M,
                'N': N,
                'h-index': h_index,
                'g-index': g_index,
                'new-g-index': new_g_index,
                'a-index': a_index,
                'r-index': r_index,
                'new-r-index': new_r_index,
                'hg-index': hg_index,
                'new-hg-index': new_hg_index,
                'new-a-index': new_a_index,
                'e-index': e_index,
                'new-e-index': new_e_index,
                'h-prime': h_prime,
                'new-h-prime': new_h_prime,
                'h2-index': h2_index,
                'F1': F1,
                'F2': F2,
                'Sum of Cites': sum_of_cites  # Store Sum of Cites
            }
            result_rows.append(result_row)

    result_df = pd.DataFrame(result_rows)
    return result_df

def plot_scatter_with_regression(x, y, xlabel, ylabel, title, doc):
    slope, intercept, r_value, p_value, std_err = linregress(x, y)

    plt.figure()
    plt.scatter(x, y, label='Data points')
    plt.plot(x, slope * x + intercept, color='red', label=f'Regression line: y={slope:.2f}x+{intercept:.2f}')
    plt.xlabel(xlabel)
    plt.ylabel(ylabel)
    plt.title(f'{title}\nRegression line with R² = {r_value**2:.2f}')
    plt.legend()
    plt.grid(True)
    
    # Save plot to a file
    plot_filename = f"{title.replace(' ', '_')}.png"
    plt.savefig(plot_filename)
    
    # Add plot to the Word document
    doc.add_paragraph(title)
    doc.add_picture(plot_filename, width=Inches(6))
    os.remove(plot_filename)  # Clean up by removing the image file

def main():
    # Hide the root window
    Tk().withdraw()

    # Prompt user to select multiple Excel files
    file_paths = filedialog.askopenfilenames(title="Select Excel files", filetypes=[("Excel files", "*.xlsx"), ("Excel files", "*.xls")])

    if not file_paths:
        print("No files selected. Exiting.")
        return

    results_df = process_excel_files(file_paths)

    output_filename = 'metrics_summary.xlsx'
    results_df.to_excel(output_filename, index=False)
    print(f"Results saved to {output_filename}")

    # Create a Word document
    doc = Document()
    doc.add_heading('Regression Plots', level=1)

    # Plotting and saving regression plots in the Word document
    plot_scatter_with_regression(results_df['g-index'], results_df['new-g-index'], 'g-index', 'new-g-index', 'g-index vs new-g-index', doc)
    plot_scatter_with_regression(results_df['r-index'], results_df['new-r-index'], 'r-index', 'new-r-index', 'r-index vs new-r-index', doc)
    plot_scatter_with_regression(results_df['a-index'], results_df['new-a-index'], 'a-index', 'new-a-index', 'a-index vs new-a-index', doc)
    plot_scatter_with_regression(results_df['hg-index'], results_df['new-hg-index'], 'hg-index', 'new-hg-index', 'hg-index vs new-hg-index', doc)
    plot_scatter_with_regression(results_df['e-index'], results_df['new-e-index'], 'e-index', 'new-e-index', 'e-index vs new-e-index', doc)
    plot_scatter_with_regression(results_df['h-prime'], results_df['new-h-prime'], 'h-prime', 'new-h-prime', 'h-prime vs new-h-prime', doc)

    # Adding h2-cubed and h-squared columns and plotting
    results_df['h2-cubed'] = results_df['h2-index'] ** 3
    results_df['h-squared'] = results_df['h-index'] ** 2
    plot_scatter_with_regression(results_df['h2-cubed'], results_df['h-squared'], '(h2)^3', 'h^2', '(h2)^3 vs h^2', doc)


    # Plotting Sum of Cites vs F1 using eqn_7
    plot_scatter_with_regression(results_df['Sum of Cites'], results_df['F1'], 'Sum of Cites', 'F1', 'Sum of Cites vs F (eqn_7_10h)', doc)
    # Plotting Sum of Cites vs F2 using eqn_8
    plot_scatter_with_regression(results_df['Sum of Cites'], results_df['F2'], 'Sum of Cites', 'F2', 'Sum of Cites vs F (eqn_8_10h)', doc)


    # Save the Word document
    doc_filename = 'regression_plots.docx'
    doc.save(doc_filename)
    print(f"Plots saved to {doc_filename}")

if __name__ == '__main__':
    main()
