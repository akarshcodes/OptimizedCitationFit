import sys
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from scipy.stats import linregress
from docx import Document
from docx.shared import Inches
import os
from PyQt5.QtWidgets import QApplication, QMainWindow, QPushButton, QVBoxLayout, QFileDialog, QLabel
from PyQt5.QtCore import Qt

def calculate_h_index(citations):
    citations_sorted = np.sort(citations)[::-1]
    h_index = np.sum(citations_sorted >= np.arange(1, len(citations_sorted) + 1))
    return h_index

def calculate_g_index_alt(citations):
    citations_sorted = np.sort(citations)[::-1]
    g_index = 0
    total_citations = 0

    for i, citation_count in enumerate(citations_sorted, start=1):
        total_citations += citation_count
        if total_citations >= i**2:
            g_index = i

    return g_index

def calculate_new_g_index(h, M):
    e = 2.71828
    new_g_index = h * (np.sqrt(np.log((4 * M) / (e * h))))
    return new_g_index

def calculate_a_index(h_index, citations):
    h_papers = citations[citations >= h_index]
    a_index = np.mean(h_papers)
    return a_index

def calculate_r_index(h_index, citations):
    h_papers = citations[citations >= h_index]
    r_index = np.sqrt(np.sum(h_papers))
    return r_index

def calculate_hg_index(h_index, g_index):
    hg_index = np.sqrt(h_index * g_index)
    return hg_index

def calculate_new_a_index(M, h):
    new_a_index = ((M * h) / (M - h)) * np.log(M / h)
    return new_a_index

def calculate_e_index(h_index, citations):
    h_papers = citations[citations >= h_index]
    excess_citations = np.sum(h_papers) - h_index**2
    e_index = np.sqrt(excess_citations)
    return e_index

def calculate_new_e_index(M, h):
    new_e_index_squared = ((M * h**2) / (M - h)) * np.log(M / h) - h**2
    new_e_index = np.sqrt(new_e_index_squared)
    return new_e_index

def calculate_new_r_index(M, h):
    new_r_index = h * np.sqrt((M / (M - h)) * np.log(M / h))
    return new_r_index

def calculate_new_hg_index(h, M):
    e = 2.71828
    new_hg_index = h * (np.log((4 * M) / (e * h)))**(1/4)
    return new_hg_index

def calculate_h_prime(h, citations):
    sorted_citations = np.sort(citations)[::-1]
    Fhead = np.sum(sorted_citations[:h])
    Ftail = np.sum(sorted_citations[h:])
    h_prime = np.sqrt(Fhead / Ftail) * h
    return h_prime

def calculate_new_h_prime(M, h, N):
    numerator = (M / (M - h)) * (np.log(M / h)) - 1
    denominator = (N / (N - h)) * (np.log(N / h)) - 1
    new_h_prime = np.sqrt(numerator / denominator)*h
    return new_h_prime

def calculate_h2_index(citations):
    citations_sorted = np.sort(citations)[::-1]
    h2_index = 0

    for i in range(1, len(citations_sorted) + 1):
        if citations_sorted[i-1] >= i**2:
            h2_index = i
        else:
            break

    return h2_index

def process_csv_files(file_paths):
    result_rows = []

    for file_path in file_paths:
        df = pd.read_csv(file_path)
        citations = df['Cites'].values

        M = np.max(citations)
        N = len(citations)
        h_index = calculate_h_index(citations)
        g_index = calculate_g_index_alt(citations)
        new_g_index = calculate_new_g_index(h_index, M)
        a_index = calculate_a_index(h_index, citations)
        r_index = calculate_r_index(h_index, citations)
        hg_index = calculate_hg_index(h_index, g_index)
        new_r_index = calculate_new_r_index(M, h_index)
        new_hg_index = calculate_new_hg_index(h_index, M)
        new_a_index = calculate_new_a_index(M, h_index)
        e_index = calculate_e_index(h_index, citations)
        new_e_index = calculate_new_e_index(M, h_index)
        h_prime = calculate_h_prime(h_index, citations)
        new_h_prime = calculate_new_h_prime(M, h_index, N)
        h2_index = calculate_h2_index(citations)

        result_row = {
            'File': file_path,
            'M': M,
            'N': N,
            'h-index': h_index,
            'g-index': g_index,
            'new-g-index': new_g_index,
            'a-index': a_index,
            'r-index': r_index,
            'new-r-index': new_r_index,
            'hg-index': hg_index,
            'new-hg-index': new_hg_index,
            'new-a-index': new_a_index,
            'e-index': e_index,
            'new-e-index': new_e_index,
            'h-prime': h_prime,
            'new-h-prime': new_h_prime,
            'h2-index': h2_index
        }
        result_rows.append(result_row)

    result_df = pd.DataFrame(result_rows)
    return result_df

def plot_scatter_with_regression(x, y, xlabel, ylabel, doc, title):
    slope, intercept, r_value, p_value, std_err = linregress(x, y)

    plt.figure()
    plt.scatter(x, y, label='Data points')
    plt.plot(x, slope * x + intercept, color='red', label=f'Regression line: y={slope:.2f}x+{intercept:.2f}')
    plt.xlabel(xlabel)
    plt.ylabel(ylabel)
    plt.title(f'Regression line with R² = {r_value**2:.2f}')
    plt.legend()
    plt.grid(True)
    plt.savefig(f'{title}.png')
    plt.close()

    doc.add_heading(title, level=2)
    doc.add_picture(f'{title}.png', width=Inches(5))
    os.remove(f'{title}.png')

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.initUI()

    def initUI(self):
        self.setWindowTitle('Drag and Drop CSV Files')
        self.setGeometry(100, 100, 600, 400)

        self.label = QLabel(self)
        self.label.setText("Drag CSV files here")
        self.label.setAlignment(Qt.AlignCenter)
        self.label.setGeometry(100, 100, 400, 50)

        self.setAcceptDrops(True)

    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            event.accept()
        else:
            event.ignore()

    def dropEvent(self, event):
        file_paths = []
        for url in event.mimeData().urls():
            file_paths.append(str(url.toLocalFile()))

        self.process_files(file_paths)

    def process_files(self, file_paths):
        results_df = process_csv_files(file_paths)
        results_df.to_excel('metrics_summary.xlsx', index=False)
        print('Results saved to metrics_summary.xlsx')

        doc = Document()
        doc.add_heading('Citation Metrics Plots', level=1)

        plot_scatter_with_regression(results_df['g-index'], results_df['new-g-index'], 'g-index', 'new-g-index', doc, 'g-index vs new-g-index')
        plot_scatter_with_regression(results_df['r-index'], results_df['new-r-index'], 'r-index', 'new-r-index', doc, 'r-index vs new-r-index')
        plot_scatter_with_regression(results_df['a-index'], results_df['new-a-index'], 'a-index', 'new-a-index', doc, 'a-index vs new-a-index')
        plot_scatter_with_regression(results_df['hg-index'], results_df['new-hg-index'], 'hg-index', 'new-hg-index', doc, 'hg-index vs new-hg-index')
        plot_scatter_with_regression(results_df['e-index'], results_df['new-e-index'], 'e-index', 'new-e-index', doc, 'e-index vs new-e-index')
        plot_scatter_with_regression(results_df['h-prime'], results_df['new-h-prime'], 'h-prime', 'new-h-prime', doc, 'h-prime vs new-h-prime')

        results_df['h2-cubed'] = results_df['h2-index']**3
        results_df['h-squared'] = results_df['h-index']**2
        plot_scatter_with_regression(results_df['h2-cubed'], results_df['h-squared'], '(h2)^3', 'h^2', doc, '(h2)^3 vs h^2')

        doc.save('metrics_plots.docx')
        print('Plots saved to metrics_plots.docx')

if __name__ == '__main__':
    app = QApplication(sys.argv)
    mainWindow = MainWindow()
    mainWindow.show()
    sys.exit(app.exec_())
